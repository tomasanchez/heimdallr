"""
Adapter for reading assignments from a file.
"""
import abc
import datetime
import os
import string
import tempfile
from io import BytesIO
from itertools import dropwhile

import docx
import fitz
import joblib
import nltk
import pandas as pd
import pytextract
from fastapi import UploadFile
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.pipeline import Pipeline
from sklearn.svm import LinearSVC
from spacy import Language

from heimdallr.adapters.text_processing import CleanTextTransformer
from heimdallr.domain.models.assignment import UNKNOWN_AUTHOR, Assignment, Topic
from heimdallr.utils import content_type
from heimdallr.utils.formatting import contains_letters_or_numbers, normalize_sentence


class AssignmentReader(abc.ABC):
    @abc.abstractmethod
    def read(self, file: UploadFile) -> Assignment:
        """
        Parses a file into a domain model.

        Args:
            file (UploadFile): The file to be parsed.
        Returns:
            Assignment: A mapped document.
        """
        raise NotImplementedError


class TopicPredictor(abc.ABC):
    @abc.abstractmethod
    def predict(self, content: str) -> Topic:
        """
        Predicts the topic of an assignment.

        Args:
            content (str): The assignment text.

        Returns:
            Topic: the corresponding topic.
        """
        raise NotImplementedError


class SpacyAssignmentReader(AssignmentReader):
    SPACY_PERSON_LABEL = "PER"

    def __init__(
        self,
        nlp: Language,
        excluded_names: list[str] | None = None,
        topic_predictor: TopicPredictor | None = None,
    ):
        """
        A PDF Reader that maps the document into a domain model.

        Args:
            nlp (Language): The Natural Language Processor.
            excluded_names (list[str] | None): A list of names to be excluded from the document.
        """
        if excluded_names is None:
            self.excluded_names = [
                "Dr",
                "Ingeniero",
                "Ing",
                "Licenciado",
                "Lic",
                "MSc",
                "Mg",
                "Profesor",
                "Prof",
                "Borré",
                "Borre",
                "Hernán Borré",
                "Hernan Borré",
                "Hernán Borre",
                "Hernan Borre",
                "Maximiliano Bracho",
                "Alejandro Bracho",
                "Alejandro Prince",
                "Prince",
                "Bracho",
                "Sistemas",
                "SISTEMAS",
            ]

        self.exclude_from_name = [
            "Dr",
            "Ingeniero",
            "Ing",
            "Licenciado",
            "Lic",
            "MSc",
            "Mg",
            "Profesor",
            "Prof",
            "Alumno",
            "Legajo",
            "Carrera",
            "Sistemas",
            "Ayudantes",
        ]

        self.nlp = nlp

        self.topic_predictor = topic_predictor

    def read(self, file: UploadFile) -> Assignment:
        assignment = Assignment(content=[], date=datetime.date.today(), author=UNKNOWN_AUTHOR, title=file.filename)

        if file.content_type == content_type.APPLICATION_PDF:
            assignment = self.read_pdf(file)

        if file.content_type == content_type.APPLICATION_DOCX:
            assignment = self.read_docx(file)

        if file.content_type == content_type.APPLICATION_WORD:
            assignment = self.read_msword(file)

        if self.topic_predictor and assignment.content:
            assignment.topic = self.topic_predictor.predict(" ".join(assignment.content))

        return assignment

    def read_pdf(self, file: UploadFile) -> Assignment:
        """
        Parses a PDF file into a domain model.

        Args:
            file: The file to be parsed.

        Returns:
            Assignment: A mapped document.
        """
        file_document = fitz.Document(stream=file.file.read(), filetype="pdf")

        pages: list[list[str]] = [self._parse_page(page.get_textpage().extractText()) for page in file_document]

        author = self._find_out_author(file_document[0].get_textpage().extractText())

        file_document.close()

        return Assignment(
            author=author,
            title=file.filename,
            content=[sentence for page in pages for sentence in page],
            date=datetime.date.today(),
        )

    def read_docx(self, file: UploadFile) -> Assignment:
        """
        Parses a DOCX file into a domain model.

        Args:
            file (UploadFile): a file reference.

        Returns:
            Assignment: A mapped document.
        """
        # obtain the file as a stream compatible with python-docx Document
        bytes_io = BytesIO()
        bytes_io.write(file.file.read())

        doc = docx.Document(bytes_io)

        paragraphs = [p.text for p in doc.paragraphs if contains_letters_or_numbers(p.text)]

        content = [sentence for paragraph in paragraphs for sentence in self._parse_page(paragraph)]

        maybe_authors = [self._find_out_author(sentence) for sentence in content]

        author = next(dropwhile(lambda name: name == UNKNOWN_AUTHOR, maybe_authors), UNKNOWN_AUTHOR)

        return Assignment(
            author=author,
            title=file.filename,
            content=content,
            date=datetime.date.today(),
        )

    def read_msword(self, file: UploadFile) -> Assignment:
        """
        Parses a .doc file into a domain model.

        Args:
            file (UploadFile): a file reference.

        Returns:
            Assignment: A mapped document.
        """

        # create a temporary file with a specific name
        with tempfile.NamedTemporaryFile(delete=False, prefix="temp_doc_", suffix=".doc") as temp_file:
            temp_file.write(file.file.read())

        text = pytextract.process(temp_file.name, extension="doc")

        content = list(self._parse_page(text)) if text else []

        maybe_authors = list(self._find_out_author(sentence) for sentence in content) if text else []

        author = next(dropwhile(lambda name: name == UNKNOWN_AUTHOR, maybe_authors), UNKNOWN_AUTHOR)

        temp_file.close()
        os.remove(temp_file.name)

        return Assignment(
            author=author,
            title=file.filename,
            content=content,
        )

    def _parse_page(self, page_text: str) -> list[str]:
        """
        Parses string from a page into a list of sentences.

        Args:
            page_text: the string to be parsed.

        Returns:
            list[str]: A list of sentences.
        """

        # Parse the page text into a spacy document
        doc = self.nlp(page_text.replace("●", ""))

        # Split the page text into sentences
        page_sentences = [normalize_sentence(str(s)) for s in doc.sents if contains_letters_or_numbers(str(s))]

        splits = [sentence.split("\n") for sentence in page_sentences]

        return [sentence for sublist in splits if sublist for sentence in sublist]

    def _find_out_author(self, page_text: str) -> str:
        """
        Finds out the author of the assignment.
        It supposes that the document contains a page with the author's name on it.

        Args:
            page_text: The first page of the assignment.

        Returns:
            str: The author's name.
        """
        page_text = page_text.replace("\n", " ")

        doc = self.nlp(page_text)

        names = [
            ent.text.title()
            for ent in doc.ents
            if (ent.label_ == self.SPACY_PERSON_LABEL and not any(name in ent.text for name in self.excluded_names))
        ]

        possible_name = next(dropwhile(lambda name: name == UNKNOWN_AUTHOR, names), None)

        if not possible_name:
            return UNKNOWN_AUTHOR

        for noun in self.exclude_from_name:
            possible_name = possible_name.replace(noun, "")

        return possible_name.strip()


class SklearnTopicPredictor(TopicPredictor):
    PRON_LABEL = "-PRON-"

    def __init__(self, nlp: Language, download: bool = False, model_path: str | None = None):
        """
        A Topic Predictor that uses a trained model to predict the topic of an assignment.

        Args:
            nlp: A Natural Language Processor
            download: Whether to download NLTK stopwords or not
            model_path: path to a .joblib model
        """
        if download:
            nltk.download("stopwords")

        self.nlp: Language = nlp
        self.stop_words: list[str] = nltk.corpus.stopwords.words("spanish")
        self.symbols = " ".join(string.punctuation).split(" ") + ["-", "...", "”", "”", "'", "“", "¿"]
        self.count_v = CountVectorizer(tokenizer=self.tokenize, ngram_range=(1, 1))
        self.clf = LinearSVC()

        # load a model
        if model_path:
            self.pipeline = joblib.load(model_path)
        else:
            self.pipeline = Pipeline(
                [("clean_up", CleanTextTransformer()), ("vectorize", self.count_v), ("clf", self.clf)]
            )

    def predict(self, content: str) -> Topic:
        """
        Predicts the topic of an assignment.

        Args:
            content (str): The assignment text.

        Returns:
            Topic: the corresponding topic.
        """
        [label] = self.pipeline.predict([content])
        return Topic(label)

    def cleanup_text(self, docs: list[str]) -> pd.Series:
        """
        Applies some pre-processing on the given list of texts.

        Args:
            docs: a list of texts to be processed.

        Returns:
            Series: The processed texts.
        """
        texts = []

        for doc in docs:
            tokens = self.tokenize(doc, disable=["parser", "ner"])
            text = " ".join([tok for tok in tokens if tok not in self.symbols])
            texts.append(text)

        return pd.Series(texts)

    def tokenize(self, sample: str, disable: list[str] | None = None) -> list[str]:
        """
        Given a sample text, apply some pre-processing and tokenize it.

        Args:
            sample: The text to be tokenized.
            disable: A list of spacy components to be disabled.

        Returns:
            list[str]: The tokenized text.
        """
        tokens = self.nlp(sample, disable=disable or [])
        lemmas = []

        for tok in tokens:
            lemmas.append(tok.lemma_.lower().strip() if tok.lemma_ != self.PRON_LABEL else tok.lower_)

        return [lemma for lemma in lemmas if lemma not in self.stop_words and lemma not in self.symbols]
